from uuid import UUID
from datetime import datetime, timezone
from io import BytesIO

from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form, Query
from sqlalchemy.orm import Session
from docx import Document

from app.db.deps import get_db
# from app.auth.deps import require_roles
from app.models import Call, CallAnalysisResult
from app.repositories.call_repository_sqlalchemy import CallRepositorySQLAlchemy
from app.repositories.call_insight_repository_sqlalchemy import CallInsightRepositorySQLAlchemy
from app.repositories.interviewee_profile_repository_sqlalchemy import IntervieweeProfileRepositorySQLAlchemy
from app.schemas.call import CallCreateRequest, CallResponse
from app.services import get_openrouter_service
from app.repositories.search_repository import SearchRepository
from app.schemas.search import AgentSearchRequest, AgentSearchResponse
from app.services.agent_call_search_service import AgentCallSearchService

# router = APIRouter(prefix="/calls", tags=["Calls"], dependencies=[Depends(require_roles(["admin", "super-admin"]))])
router = APIRouter(prefix="/calls", tags=["Calls"])



@router.post("/", response_model=CallResponse, status_code=status.HTTP_201_CREATED)
def create_call(payload: CallCreateRequest, db: Session = Depends(get_db)):
    repo = CallRepositorySQLAlchemy(db)
    call = Call(
        title=payload.title,
        description=payload.description,
        call_date=payload.call_date,
        transcript=payload.transcript,
    )

    saved_call = repo.create(call)
    return CallResponse(
        id=saved_call.id,
        title=saved_call.title,
        description=saved_call.description,
        call_date=saved_call.call_date,
        status=saved_call.status,
        created_at=saved_call.created_at,
    )


@router.get("/", response_model=list[CallResponse])
def list_calls(
    status_filter: str | None = Query(None, alias="status"),
    limit: int = Query(50, ge=1, le=200),
    offset: int = Query(0, ge=0),
    db: Session = Depends(get_db),
):
    """List calls with optional filters and pagination."""
    repo = CallRepositorySQLAlchemy(db)
    calls = repo.list(status=status_filter, limit=limit, offset=offset)
    return [
        CallResponse(
            id=c.id,
            title=c.title,
            description=c.description,
            call_date=c.call_date,
            status=c.status,
            created_at=c.created_at,
        )
        for c in calls
    ]


@router.post("/upload-docx", response_model=CallResponse, status_code=status.HTTP_201_CREATED)
async def create_call_from_docx(
    title: str | None = Form(None),
    description: str | None = Form(None),
    call_date: datetime | None = Form(None),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    """Create a call by uploading a meeting transcript in .docx format."""
    filename = (file.filename or "").lower()
    if not filename.endswith(".docx"):
        raise HTTPException(status_code=415, detail="Only .docx files are supported")

    raw = await file.read()
    if not raw:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")

    try:
        doc = Document(BytesIO(raw))
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid .docx file")

    parts: list[str] = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                ct = (cell.text or "").strip()
                if ct:
                    cells.append(ct)
            if cells:
                parts.append(" | ".join(cells))

    transcript = "\n".join(parts).strip()

    if len(transcript) < 50:
        raise HTTPException(
            status_code=422,
            detail="Transcript extracted from DOCX is too short (min 50 characters).",
        )

    if call_date is None:
        call_date = datetime.now(timezone.utc)

    if not title:
        # Default title: use filename (without extension) or date
        base = (file.filename or "").rsplit(".", 1)[0].strip()
        title = base or f"Call on {call_date.date().isoformat()}"

    repo = CallRepositorySQLAlchemy(db)
    call = Call(
        title=title,
        description=description,
        call_date=call_date,
        transcript=transcript,
    )
    saved_call = repo.create(call)

    return CallResponse(
        id=saved_call.id,
        title=saved_call.title,
        description=saved_call.description,
        call_date=saved_call.call_date,
        status=saved_call.status,
        created_at=saved_call.created_at,
    )


@router.post("/{call_id}/analyze", response_model=CallAnalysisResult, status_code=status.HTTP_200_OK)
async def analyze_call(call_id: UUID, db: Session = Depends(get_db)):
    calls_repo = CallRepositorySQLAlchemy(db)
    insight_repo = CallInsightRepositorySQLAlchemy(db)
    interviewee_profile_repo = IntervieweeProfileRepositorySQLAlchemy(db)

    call = calls_repo.get(call_id)
    if not call:
        raise HTTPException(status_code=404, detail="Call not found")

    service = get_openrouter_service()
    if not service.is_configured:
        raise HTTPException(
            status_code=503,
            detail="OpenRouter is not configured (missing OPENROUTER_API_KEY)",
        )

    calls_repo.update_status(call_id, "processing")
    analysis_result = await service.analyze_call(call.transcript)

    if not analysis_result:
        calls_repo.update_status(call_id, "failed")
        raise HTTPException(status_code=502, detail="Call analysis failed")

    call_searchable_text = build_call_searchable_text(
        call_title=call.title,
        call_description=call.description,
        insights=analysis_result.insights,
    )
    call_embedding = await service.generate_embedding(call_searchable_text)
    call_insight = insight_repo.upsert(
        call_id,
        analysis_result.insights,
        searchable_text=call_searchable_text,
        embedding=call_embedding,
    )
    if analysis_result.interviewee_profile and analysis_result.interviewee_profile is not {} and analysis_result.insights.call_type == "interview":
        searchable_text = build_searchable_summary(analysis_result.interviewee_profile.to_dict())
        embedding = await service.generate_embedding(searchable_text)
        analysis_result.interviewee_profile.searchable_summary = f"{searchable_text} \n {analysis_result.interviewee_profile.searchable_summary or ''}"
        analysis_result.interviewee_profile.embedding = embedding
        analysis_result.interviewee_profile.has_pe_experience = "private_equity_exposure" in analysis_result.interviewee_profile.to_dict() and bool(analysis_result.interviewee_profile.to_dict().get("private_equity_exposure"))
        interviewee_profile_saved =  interviewee_profile_repo.upsert(call_insight_id=call_insight["id"], profile=analysis_result.interviewee_profile)
        analysis_result.interviewee_profile.id = interviewee_profile_saved["id"]
    calls_repo.update_status(call_id, "analyzed")
    analysis_result.call_id = call_id
    analysis_result.insights.id = call_insight["id"]
    return analysis_result


@router.get("/{call_id}/insights", response_model=CallAnalysisResult, status_code=status.HTTP_200_OK)
def get_call_insights(call_id: UUID, db: Session = Depends(get_db)):
    calls_repo = CallRepositorySQLAlchemy(db)
    insight_repo = CallInsightRepositorySQLAlchemy(db)

    call = calls_repo.get(call_id)
    if not call:
        raise HTTPException(status_code=404, detail="Call not found")

    insights = insight_repo.get(call_id)
    interviewee_profile = None

    if not insights:
        raise HTTPException(status_code=404, detail="No insights found for this call")
    
    if insights.call_type == "interview":
        interviewee_profile_repo = IntervieweeProfileRepositorySQLAlchemy(db)
        interviewee_profile = interviewee_profile_repo.get(insights.id)

    return CallAnalysisResult(call_id=call_id, insights=insights, interviewee_profile=interviewee_profile, )

def build_searchable_summary(profile: dict) -> str:
    return f"""
    Name: {profile.get('full_name')}
    Title: {profile.get('current_title')}
    Company: {profile.get('current_company')}
    Seniority: {profile.get('seniority_level')}

    Industries: {profile.get('industry_focus')}
    Transformations: {profile.get('transformation_experience')}
    PE Exposure: {profile.get('private_equity_exposure')}

    Leadership Scope: {profile.get('leadership_scope')}
    Achievements: {profile.get('notable_achievements')}
    """


def build_call_searchable_text(*, call_title: str, call_description: str | None, insights: CallAnalysisResult | None = None, ):
    """Build a compact text blob to embed for call-level semantic search."""
    if insights is None:
        return f"Title: {call_title}\nDescription: {call_description or ''}".strip()

    meta = insights.insights if hasattr(insights, "insights") else insights
    # meta is CallInsight
    tags = ", ".join(getattr(meta, "tags", []) or [])
    key_decisions = " | ".join(getattr(meta, "key_decisions", []) or [])
    people = ", ".join([p.name for p in (getattr(meta, "people_mentioned", []) or []) if getattr(p, "name", None)])
    action_items = " | ".join([ai.description for ai in (getattr(meta, "action_items", []) or []) if getattr(ai, "description", None)])

    parts = [
        f"Title: {call_title}",
        f"Description: {call_description or ''}",
        f"Call Type: {getattr(meta, 'call_type', '')}",
        f"Summary: {getattr(meta, 'summary', '')}",
    ]
    if tags:
        parts.append(f"Tags: {tags}")
    if people:
        parts.append(f"People Mentioned: {people}")
    if key_decisions:
        parts.append(f"Key Decisions: {key_decisions}")
    if action_items:
        parts.append(f"Action Items: {action_items}")

    return "\n".join([p for p in parts if p]).strip()

@router.post("/search", response_model=AgentSearchResponse)
async def agent_search_calls(request: AgentSearchRequest, db: Session = Depends(get_db)):
    """AI-agent style search over calls and interviewee profiles.

    Supports:
    - retrieval: semantic search over call summaries and interviewee profiles
    - comparative: semantic retrieval + LLM scoring/ranking
    - aggregation: database stats (counts)
    """
    service = get_openrouter_service()
    if not service.is_configured:
        raise HTTPException(
            status_code=503,
            detail="OpenRouter is not configured (missing OPENROUTER_API_KEY)",
        )

    try:
        agent = AgentCallSearchService(openrouter_service=service, db=db)
        return await agent.search(
            query=request.query,
            top_k=request.top_k,
            similarity_threshold=request.similarity_threshold,
        )
    except Exception as e:
        print(f"Error during agent search: {e}")
        raise HTTPException(status_code=500, detail=str(e))